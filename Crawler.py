import logging
import os
import datetime
import re
import json

from Article import Article
from ContentCrawler import *
from WordProcess import *

import urllib
import requests
import mechanicalsoup
from bs4 import BeautifulSoup
from urllib.parse import quote

import docx
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from docx.enum.text import *
from docx.enum.table import *
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls,qn
from docx.text.paragraph import Paragraph

import smtplib
from email.header import Header
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

with open('app.json','r', encoding='UTF-8') as json_file:
    config = json.load(json_file)

generalInclude = (config['ntcInclude'], config['memInclude'], config['semiInclude'], config['indInclude'])
smgInclude = config['smgInclude']
hynixInclude = config['hynixInclude']
micronInclude = config ['micronInclude']
exclude = config['exclude']
articleExclude = config['articleExclude']

articles=[]
results = [[] for i in range(4)]
digiTimesList = []
imgNum = 0
archive = []
headers = {'User-Agent': 'Mozilla/5.0'}

def basicCrawler(name, url, parentType, parentClass, childType, childClass, titleType, titleClass, hrefType, isRelHref, newsCount = 100):
  try:
    #mainSite = ''
    href = ''
    title=''
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    #logging.warning(soup)
    count = 0
    for ele in soup.find(parentType, parentClass).find_all(childType, childClass):
        if count >= newsCount:
          break
        #logging.warning(ele)
        if ele.find(hrefType):
          if titleClass == '':
            if ele.find(titleType):
              title = ele.find(titleType).text
              #logging.warning(ele)
          else:
            if ele.find(titleType,titleClass):
              title = ele.find(titleType,titleClass).text
              #logging.warning(ele)
          href = ele.find(hrefType).get('href')

          if href.startswith('//'):
            href = 'https:' + href
          elif href.startswith('http') == False: 
            href = url[:url.find('/', 8)] + href
          
          #logging.warning(title)
          #logging.warning(mainSite + ele.find(hrefType).get('href'))
          if title != '':
            articles.append(Article.from_list(title, href, name))
            count += 1
  except Exception as e:
    logging.warning(url + ':' + str(e))

def formattedCrawler(name, url, contentType, contentClass, titleType, titleClass, hrefType, isRelHref):
  try:
    title = ''
    mainSite = ''

    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    #logging.warning(soup)
    for ele in soup.find_all(contentType, contentClass):
        #logging.warning(ele.text)
        # if titleClass == '':
        #     title = ele.find(titleType).text
        # else:
        #     title = ele.find(titleType,titleClass).text
        title = ele.text
        #logging.warning(ele.parent.get('href'))

        if isRelHref:
            mainSite = url[:url.find('/', 8)]
            #logging.warning(mainSite)
        else:
            mainSite = ''
        
        url = mainSite + ele.parent.get('href')
        url = url.replace('\n','').replace('\r','')
        articles.append(Article.from_list(title, url, name))
  except Exception as e:
    logging.warning(name + ':' + str(e))

def articleCrawler(name, url, contentType,titleType, titleClass, hrefType, isRelHref):
  try:
    mainSite = ''
    title=''
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    #logging.warning(soup)
    for ele in soup.find_all(contentType):
        #logging.warning(ele)
        if titleClass == '':
          title = ele.find(titleType).text
        else:
          title = ele.find(titleType,titleClass).text
          
        #logging.warning(title)
        if isRelHref:
            mainSite = url[:url.find('/', 8)]
        else:
            mainSite = ''

        articles.append(Article.from_list(title, mainSite + ele.find(hrefType).get('href'), name))
  except Exception as e:
    logging.warning(name + ':' + str(e))

def saveImage(article):
  #logging.warning(img_url)
  try:
    global imgNum
    # #create the object, assign it to a variable
    # proxy = urllib.request.ProxyHandler(proxies)
    # # construct a new opener using your proxy settings
    # opener = urllib.request.build_opener(proxy)
    # # install the openen on the module-level
    # urllib.request.install_opener(opener) 
    # #if title.endswith('jpg') == False and title.endswith('png') == False and title.endswith('gif') == False:
    # #  title += '.jpg'

    #article.img_url = quote(article.img_url,safe='/:?=')
    urllib.request.urlretrieve(quote(article.img_url,safe='/:?='), os.path.join('Temp', 'img_' + str(imgNum) + '.jpg'))
    article.img_name = 'img_' + str(imgNum) + '.jpg'
    imgNum += 1
  except Exception as e:
    logging.warning('Save Image Error:"' + article.img_url + '"' + str(e))

def filterArticles(digiTimes = False):
  resultList = []
  firstRun = 'DIGITIMES電子時報'

  if digiTimes:
    firstRun = ''

  newsList = []
  

  try:
    with open('News List.txt', 'r') as f2:
      line = f2.readline()
      while line is not None and line != '':
        #logging.warning(line.strip())
        line = f2.readline()
        newsList.append(line.strip())
  except Exception as e:
    logging.warning(e)

  for article in articles:
    if article.title.strip() not in resultList:
      if article.url.strip() in newsList:
        archive.append(article.url.strip())
      elif article.source != firstRun:
        print(article.title.strip())
        
        #group = 1

        if article.source == 'Samsung Pressroom':
          include = smgInclude
        elif article.source == 'SK Hynix Pressroom':
          include = hynixInclude
        elif article.source == 'Micron Pressroom':
          include = micronInclude
        # else:
        #   includeList = generalInclude
        
        match = False
        if article.source == 'Samsung Pressroom' or article.source == 'SK Hynix Pressroom' or article.source == 'Micron Pressroom':
          i = 0
          while match == False and i < len(include):
            if type(include[i]) == list:
              if all(s.lower() in article.title.lower() for s in include[i]):
                results[1].append(article)
                resultList.append(article.title.strip())
                match = True
            elif include[i].lower() in article.title.lower():
              results[1].append(article)
              resultList.append(article.title.strip())
              match = True
              
            i += 1
        else:
          for group in range(4):
            i = 0
            include = generalInclude[group]
            #logging.warning(include[0])
            while match == False and i < len(include):
              if type(include[i]) == list:
                if all(s.lower() in article.title.lower() for s in include[i]):
                  results[group].append(article)
                  resultList.append(article.title.strip())
                  match = True
              elif include[i].lower() in article.title.lower():
                results[group].append(article)
                resultList.append(article.title.strip())
                match = True
                
              i += 1
  
  for result in results:
    for article in result:
      #print(article.title)
      match = False
      if article.source != firstRun:
        i = 0
        while match == False and i < len(exclude):
          if type(exclude[i]) == list:
            if all(s.lower() in article.title.lower() for s in exclude[i]):
              result.remove(article)
              match = True
              #print(exclude[i])
          elif exclude[i].lower() in article.title.lower():
            result.remove(article)
            match = True
            #print(exclude[i])
                   
          i += 1
          
        y = 0
        while match == False and y < len(articleExclude):
          if type(articleExclude[y]) == list:
            if all(s.lower() in article.text.lower() for s in articleExclude[y]):
              result.remove(article)
              match = True
              #print(articleExclude[y])
          elif articleExclude[y].lower() in article.text.lower():
            result.remove(article)
            match = True
            #print(articleExclude[y])
          
          y += 1
          
def loadExternalNews():
  try:
    with open('External.txt', 'r',encoding = 'utf-8') as json_file:
        array = json.load(json_file)
    
    global articles
    for article in array:
      articles.append(Article.from_external(article['url'], article['title'], article['source'], article['author'], article['text']))
      #articles.append(Article.from_external(array['title,url,'DigiTimes'))
    #articles += array
    #print(len(array))
  except Exception as e:
    logging.warning('Load External News Error:' + str(e))
  
def contentCrawlerView(article):
  try:
      if article.source == '鉅亨網':
        directInfoCrawler(article,'time','', 'span', '_3lKe')
        basicContentCrawler(article,'div','_2E8y','p')
      elif article.source == 'TechNews':
        publishInfoCrawler(article,'span', 'body', 'a')
        #basicContentCrawler(article,'div','indent','p')

        text = ''
        dom = requests.get(article.url,  verify = False).text
        soup = BeautifulSoup(dom, 'lxml')
        article_body = soup.find('div', 'indent')
        img_body = soup.find('div', 'bigg')
        #logging.warning(soup)

        for ele in article_body.find_all('p'):
          text += ele.text + '\n'

        article.text = text.strip()
        if img_body.find('img'):
          article.img_url = img_body.find('img').get('src')

          if article.img_url.startswith('//'):
            article.img_url = 'https:' + article.img_url
          elif article.img_url.startswith('http') == False: 
            article.img_url = article.url[:article.url.find('/', 8)] + article.img_url
      
      elif article.source == '日經':
        text = ''
        dom = requests.get(article.url,  verify = False).text
        soup = BeautifulSoup(dom, 'lxml')
        article_body = soup.find('div', 'newsText fix')
        #logging.warning(soup)

        for ele in article_body.find_all('p'):
          text += ele.text + '\n'

        #article.text += text
        directInfoCrawler(article,'p','time','','')
        nextPage = soup.find('span','article_pagination_next')

        while(nextPage):
          dom = requests.get(nextPage.find('a').get('href'),  verify = False).text
          soup = BeautifulSoup(dom, 'lxml')
          article_body = soup.find('div', 'newsText fix')
          #logging.warning(soup)

          for ele in article_body.find_all('p'):
            text += ele.text + '\n'
          
          nextPage = soup.find('span','article_pagination_next')

        article.text += text.strip()

        if article_body.find('img'):
          article.img_url = article_body.find('img').get('src')

          if article.img_url.startswith('//'):
            article.img_url = 'https:' + article.img_url
          elif article.img_url.startswith('http') == False: 
            article.img_url = article.url[:article.url.find('/', 8)] + article.img_url
          #logging.warning(article.img_url)
      elif article.source == 'Taipei Times':
          basicContentCrawler(article,'div','text','p')
          directInfoCrawler(article, 'div', 'date', 'div', 'reporter')
      elif article.source == 'BusinessKorea':
        basicContentCrawler(article,'div','user-snb-wrapper','p')
        listInfoCrawler(article,'div','info-text','li','')
      elif article.source == '韓聯社':
        basicContentCrawler(article,'div','article-story','p')
        publishInfoCrawler(article,'div', 'info-con', 'span')
      elif article.source == '騰訊新聞':
        basicContentCrawler(article,'div','content-article','p')
        publishInfoCrawler(article,'div', 'info-con', 'span')
      elif article.source == 'PTT':
        text = ''
        dom = requests.get(article.url,  verify = False).text
        soup = BeautifulSoup(dom, 'lxml')
        article_body = soup.find('div',id='main-content')


        article.text = article_body.text.strip()
        #logging.warning(article.title)
      elif article.source == '經濟日報':
        text = ''
        dom = requests.get(article.url,  verify = False, headers=headers).text
        soup = BeautifulSoup(dom, 'lxml')
        article_body = soup.find('div',id='article_body')

        for ele in article_body.find_all('p'):
          text += ele.text + '\n'

        article.text = text.strip()
        #logging.warning(article.title)
        
        if article_body.find('img'):
          article.img_url = article_body.find('img').get('src')

          if article.img_url.startswith('//'):
            article.img_url = 'https:' + article.img_url
          elif article.img_url.startswith('http') == False: 
            article.img_url = article.url[:article.url.find('/', 8)] + article.img_url
          #logging.warning(article.img_url)  

        directInfoCrawler(article,'div','shareBar__info--author','span','')
      elif article.source == '中時電子報':
        basicContentCrawler(article,'div','article-body','p')
        directInfoCrawler(article,'time', '', 'div', 'author')
      elif article.source == '自由時報': #解不了
        text = ''
        dom = requests.get(article.url,  verify = False).text
        soup = BeautifulSoup(dom, 'lxml')
        article_body = soup.find('div',['text','news_content'])

        for ele in article_body.find_all('p'):
          text += ele.text + '\n'

        article.text = text.strip()

        if article_body.find('img'):
          #logging.warning(article_body.find('img'))
          article.img_url = article_body.find('img').get('src')

          if article.img_url.startswith('//'):
            article.img_url = 'https:' + article.img_url
          elif article.img_url.startswith('http') == False: 
            article.img_url = article.url[:article.url.find('/', 8)] + article.img_url
          #logging.warning(article.img_url)

        article.date = article_body.find('span').text.strip()
      elif article.source == 'MoneyDJ':
        #basicContentCrawler(article,'article', '', 'p')

        try:
          text = ''
          dom = requests.get(article.url,  verify = False).text
          soup = BeautifulSoup(dom, 'lxml')
          article.date = soup.find('span', id='MainContent_Contents_lbDate').text.strip()
          #logging.warning(article.date)
          article_body = soup.find('article')
          #logging.warning(soup)

          text = article_body.text

          #for ele in article_body.find_all({'p','a'}):
          #  text += ele.text + '\n'

          article.text = text.strip()
          if article_body.find('img'):
            article.img_url = article_body.find('img').get('src')

            if article.img_url.startswith('//'):
              article.img_url = 'https:' + article.img_url
            elif article.img_url.startswith('http') == False: 
              article.img_url = article.url[:article.url.find('/', 8)] + article.img_url
            #logging.warning(article.img_url)
        except Exception as e:
          logging.warning(e)

        #dom = requests.get(article.url,  verify = False).text
        #soup = BeautifulSoup(dom, 'lxml')
        article.date = soup.find('span', id='MainContent_Contents_lbDate').text.strip()
        #logging.warning(article.date)
      elif article.source == 'ETToday':
        basicContentCrawler(article,'div', 'story', 'p')
        directInfoCrawler(article,'time', 'date', '', '')
      elif article.source == 'Yahoo奇摩股市':
        #basicContentCrawler(article,'td', 'yui-text-left', 'p')
        try:
          text = ''
          dom = requests.get(article.url,  verify = False).text
          soup = BeautifulSoup(dom, 'lxml')
          article_body = soup.find('td', 'yui-text-left')
          
          info = article_body.find_all('span','t1')
          article.date = info[0].text.strip()
          article.author = info[1].text.strip()

          for ele in article_body.find_all('p'):
            text += ele.text + '\n'

          article.text = text.strip()
          if article_body.find('img'):
            article.img_url = article_body.find('img').get('src')

            if article.img_url.startswith('//'):
              article.img_url = 'https:' + article.img_url
            elif article.img_url.startswith('http') == False: 
              article.img_url = article.url[:article.url.find('/', 8)] + article.img_url
            #logging.warning(article.img_url)
        except Exception as e:
          logging.warning(e)
      elif article.source == '中央社':
        basicContentCrawler(article,'div', 'paragraph', 'p')
      elif article.source == '數位時代':
        basicContentCrawler(article,'article', 'main_content', 'p')

        directInfoCrawler(article, 'span','item','date','author')
      elif article.source == '非凡':
        text = ''
        dom = requests.get(article.url,  verify = False).text
        soup = BeautifulSoup(dom, 'lxml')

        article.date = soup.find('h1','module-title h3 title').find('div').text.strip()
        #logging.warning(article.url)
        #logging.warning(soup)

        for ele in soup.find_all('p','main-content story'):
          text += ele.text + '\n'

        article.text = text.strip()
      elif article.source == '財訊科技':
        basicContentCrawler(article,'article', 'b-post b-post-full clearfix', 'p')
      elif article.source == 'EE Times':
        basicContentCrawler(article,'div','articleBody','p')
      elif article.source == 'SEMI Pressroom':
        basicContentCrawler(article,'div', 'field-item even', 'p')
      elif article.source == 'Trend Force Pressroom':
        basicContentCrawler(article,'div', 'content', 'p')
        directInfoCrawler(article,'span','con-date', 'span', 'analyser')
      elif article.source == 'SEMI Pressroom US':
        basicContentCrawler(article,'div', 'press-release__body', 'p')
      elif article.source == 'SIA':
        basicContentCrawler(article,'main', 'site-main', 'p')

        dom = requests.get(article.url,  verify = False).text
        soup = BeautifulSoup(dom, 'lxml')
        soup.find('main', 'site-main').find('h1').extract()
        article.date = soup.find('main', 'site-main').text.strip()
        #logging.warning(article.date)
      elif article.source == '蘋果日報':
        basicContentCrawler(article,'div', 'ndArticle_margin', 'p')
        directInfoCrawler(article, 'div','ndArticle_creat','','')
      elif article.source == 'CNBC':
        basicContentCrawler(article,'div', 'main__right', 'p')
        #directInfoCrawler(article, 'div','ndArticle_creat','','')
      elif article.source == '新浪香港':
        basicContentCrawler(article,'div', 'news-body', 'p')
        directInfoCrawler(article,'div','news-datetime', '', '')
      elif article.source == '多維新聞網':
        #basicContentCrawler(article,'div', 'dia-lead-one', 'p')
        #directInfoCrawler(article,'div','r', '', '')

        try:
          text = ''
          dom = requests.get(article.url,  verify = False)
          dom = dom.content.decode('utf-8')
          soup = BeautifulSoup(dom, 'lxml')

          if soup.find('div','r'):
            article.date = soup.find('div','r').text.strip()

          article_body = soup.find('div', 'dia-lead-one')
          #logging.warning(soup)

          for ele in article_body.find_all('p'):
            text += ele.text + '\n'

          article.text = text.strip()
          if article_body.find('img'):
            article.img_url = article_body.find('img').get('src')

            if article.img_url.startswith('//'):
              article.img_url = 'https:' + article.img_url
            elif article.img_url.startswith('http') == False: 
              article.img_url = article.url[:article.url.find('/', 8)] + article.img_url
            #logging.warning(article.img_url)
        except Exception as e:
          logging.warning(article.url + '  Basic Content Error:' + str(e))
      elif article.source == 'TechOrange':
        try:
          text = ''
          dom = requests.get(article.url,  verify = False)
          dom = dom.content.decode('utf-8')
          soup = BeautifulSoup(dom, 'lxml')

          if soup.find('time','entry-date published'):
            article.date = soup.find('time','entry-date published').text.strip()

          if soup.find('span', 'author vcard'):
            article.author = soup.find('span', 'author vcard').text.strip()

          article_body = soup.find('div', 'entry-content')
          #logging.warning(soup)

          for ele in article_body.find_all('p'):
            text += ele.text + '\n'

          article.text = text.strip()
          if article_body.find('img'):
            article.img_url = article_body.find('img').get('src')

            if article.img_url.startswith('//'):
              article.img_url = 'https:' + article.img_url
            elif article.img_url.startswith('http') == False: 
              article.img_url = article.url[:article.url.find('/', 8)] + article.img_url
            #logging.warning(article.img_url)
        except Exception as e:
          logging.warning(article.url + '  Basic Content Error:' + str(e))

        # basicContentCrawler(article,'div', 'entry-content', 'p')
        # directInfoCrawler(article,'time','entry-date published', 'span', 'author vcard')
      elif article.source != 'DIGITIMES電子時報':
        newsPleaseContentCrawler(article)

        if article.source == 'WSTS':
          article.img_url=''
          directInfoCrawler(article,'div','date', '', '')
        elif article.source == 'IC Insights Pressroom':
          article.img_url=''
          directInfoCrawler(article,'strong','midtitle2', '', '')
        elif article.source == 'Gartner Pressroom':
          directInfoCrawler(article,'h5','grid-norm', '', '')
        elif article.source == 'Samsung Pressroom':
          directInfoCrawler(article,'div','meta', '', '')
        elif article.source == 'Micron Pressroom':
          directInfoCrawler(article,'div','field field--name-field-nir-news-date field--type-datetimezone field--label-hidden', '', '')

      if article.img_url is not None and article.img_url != '' and article.source != 'DIGITIMES電子時報':
        saveImage(article)
  except Exception as e:
    logging.warning(article.url + ' Content Crawler Error:' + str(e))

def digiTimesLogin():
  browser = mechanicalsoup.Browser()
  login_page = browser.get("https://www.digitimes.com.tw/tech/lgn/lgn.asp?tourl=/tech/default.asp?",  verify = False)
  login_form = login_page.soup.find("form",id='Login')
  #logging.warning(login_form)

  login_form.find("input", {"name": "checkid"})["value"] = "zehui@ntc.com.tw"
  login_form.find("input", {"name": "checkpwd"})["value"] = "Nanya123"
  login_response = browser.submit(login_form, login_page.url,  verify = False)
  return browser

def digiTimesView():
  try:
    browser = digiTimesLogin()
    digiTimes(browser, "https://www.digitimes.com.tw/tech/")
    digiTimes(browser, "https://www.digitimes.com.tw/tech/dt/category.asp?CnlID=1")
    digiTimes(browser, "https://www.digitimes.com.tw/tech/dt/SubChannel_1_40.asp?CnlID=1&cat=40")
    digiTimes(browser, "https://www.digitimes.com.tw/tech/dt/SubChannel_1_50.asp?CnlID=1&cat=50")
    digiTimes(browser, "https://www.digitimes.com.tw/tech/dt/SubChannel_1_70.asp?CnlID=1&cat=70")

    filterArticles(True)
    
    for result in results:
      for article in result:
        getDigiTimesArticle(browser,article)
  except Exception as e:
    logging.warning('DigiTimes Error:' + str(e))

def digiTimes(browser, url):
  dom = browser.get(url,  verify = False).text
  soup = BeautifulSoup(dom, 'lxml')
  for ele in soup.find('div','wsn_wrapper_list').find_all('p'):
    text = ''
    url = 'https://www.digitimes.com.tw' + ele.find('a').get('href')
    #logging.warning(url)
    artid = re.search(r'[\?&]id=.+',url).group(0)[1:]
    #logging.warning(artid)
    if artid.find('&') != -1:
      artid = artid[:artid.find('&')]
    
    url = 'https://www.digitimes.com.tw/tech/dt/n/shwnws.asp?' + artid
    #logging.warning(url)  
    #img_url = ''
    date = ''
    title = ele.find('a').text
    global digiTimesList
    if not url.strip() in digiTimesList: 
      digiTimesList.append(url.strip())
      #logging.warning(title.strip())
      articles.append(Article.from_list(title,url,'DIGITIMES電子時報'))

def getDigiTimesArticle(browser,article):
    text = ''
    img_url=''
    date=''
    author = ''
    dom = browser.get(article.url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')

    # if soup.find('time'):
    #   date = soup.find('time').text
    #   #logging.warning(article.date)
    
    if soup.find('ul','list-inline m-b-5 txt-16'):
      author = soup.find('ul','list-inline m-b-5 txt-16').text.replace('\r','').replace('\n','\t')

    article_body = soup.find('div', id='newsText')
    #logging.warning(soup)

    for ele in article_body.find_all('p'):#,text=True, recursive=False):
        if ele.find('script'):
          ele.find('script').extract()
      #if 'function FC_ExportReady(DOMId)' in ele.text == False:
        text += ele.text + '\n'

    #article.text = text
    if article_body.find('img'):
      img_url = article_body.find('img').get('src')

      if img_url.startswith('//'):
        img_url = 'https:' + img_url
      elif img_url.startswith('http') == False: 
        img_url = img_url[:img_url.find('/', 8)] + img_url
        #logging.warning(img_url)

    article.text = text.replace('點擊圖片放大觀看','').strip()
    article.img_url = img_url
    #article.date = date.strip()
    article.author = author.strip()
    #articles.append(Article.from_digiTimes(title, url,'DigiTimes', text, date, img_url))
    if img_url != '':
      saveImage(article)

def getNewsList():
  for i in range(1,3):
    articleCrawler('TechNews','https://technews.tw/page/' + str(i) + '/','article', 'a', '', 'a', False)
  
  try:
    for i in range(1,10):
      dom = requests.get('https://money.udn.com/money/breaknews/1001/0/' + str(i),  verify = False, headers=headers).text
      soup = BeautifulSoup(dom, 'html.parser')
      for ele in soup.find('table', id='ranking_table').find_all('tr'):
        if ele.find('a'):
          articles.append(Article.from_list(ele.find('a').text, ele.find('a').get('href'),'經濟日報'))
    
    for i in range(1,10):
      dom = requests.get('https://money.udn.com/rank/newest/1001/5591/' + str(i),  verify = False, headers=headers).text
      soup = BeautifulSoup(dom, 'html.parser')
      for ele in soup.find('table', id='ranking_table').find_all('tr'):
        if ele.find('a'):
          articles.append(Article.from_list(ele.find('a').text, ele.find('a').get('href'),'經濟日報'))

    for i in range(1,10):
      dom = requests.get('https://money.udn.com/rank/newest/1001/5590/' + str(i),  verify = False, headers=headers).text
      soup = BeautifulSoup(dom, 'html.parser')
      for ele in soup.find('table', id='ranking_table').find_all('tr'):
        if ele.find('a'):
          articles.append(Article.from_list(ele.find('a').text, ele.find('a').get('href'),'經濟日報'))

    for i in range(1,10):
      dom = requests.get('https://money.udn.com/rank/pv/1001/5588/' + str(i),  verify = False, headers=headers).text
      soup = BeautifulSoup(dom, 'html.parser')
      for ele in soup.find('table', id='ranking_table').find_all('tr'):
        if ele.find('a'):
          articles.append(Article.from_list(ele.find('a').text, ele.find('a').get('href'),'經濟日報'))

    for i in range(1,10):
      dom = requests.get('https://money.udn.com/money/breaknews/1001/5589/' + str(i),  verify = False, headers=headers).text
      soup = BeautifulSoup(dom, 'html.parser')
      for ele in soup.find('table', id='ranking_table').find_all('tr'):
        if ele.find('a'):
          articles.append(Article.from_list(ele.find('a').text, ele.find('a').get('href'),'經濟日報'))
  except Exception as e:
    logging.warning('UDN' + str(e))

  try:
    for i in range(1,10):
      dom = requests.get('https://udn.com/news/breaknews/' + str(i),  verify = False, headers=headers).text
      soup = BeautifulSoup(dom, 'html.parser')
      for ele in soup.find('table', id='ranking_table').find_all('tr'):
        if ele.find('a'):
          articles.append(Article.from_list(ele.find('a').text, ele.find('a').get('href'),'經濟日報'))
          
    for i in range(1,10):
      dom = requests.get('https://udn.com/news/cate/2/6644/' + str(i),  verify = False, headers=headers).text
      soup = BeautifulSoup(dom, 'html.parser')
      for ele in soup.find('table', id='ranking_table').find_all('tr'):
        if ele.find('a'):
          articles.append(Article.from_list(ele.find('a').text, ele.find('a').get('href'),'經濟日報'))
          
    for i in range(1,10):
      dom = requests.get('https://udn.com/news/cate/2/6645/' + str(i),  verify = False, headers=headers).text
      soup = BeautifulSoup(dom, 'html.parser')
      for ele in soup.find('table', id='ranking_table').find_all('tr'):
        if ele.find('a'):
          articles.append(Article.from_list(ele.find('a').text, ele.find('a').get('href'),'經濟日報'))
          
    for i in range(1,10):
      dom = requests.get('https://udn.com/news/cate/2/6640/' + str(i),  verify = False, headers=headers).text
      soup = BeautifulSoup(dom, 'html.parser')
      for ele in soup.find('table', id='ranking_table').find_all('tr'):
        if ele.find('a'):
          articles.append(Article.from_list(ele.find('a').text, ele.find('a').get('href'),'經濟日報'))
          
    for i in range(1,10):
      dom = requests.get('https://udn.com/news/cate/2/7226/' + str(i),  verify = False, headers=headers).text
      soup = BeautifulSoup(dom, 'html.parser')
      for ele in soup.find('table', id='ranking_table').find_all('tr'):
        if ele.find('a'):
          articles.append(Article.from_list(ele.find('a').text, ele.find('a').get('href'),'經濟日報'))
         
    #惰性載入
    dom = requests.get('https://udn.com/rank/newest/2/0',  verify = False, headers=headers).text
    soup = BeautifulSoup(dom, 'html.parser')
    for ele in soup.find('table', id='ranking_table').find_all('tr'):
      if ele.find('a'):
        articles.append(Article.from_list(ele.find('a').text, ele.find('a').get('href'),'經濟日報'))
  except Exception as e:
    logging.warning('UDN' + str(e))

  for i in range(0,20,10):
    basicCrawler('日經', 'https://zh.cn.nikkei.com/industry.html?start=' + str(i),'dl', 'newsContent02', 'dt', '', 'a', '', 'a', False)
    
  basicCrawler('日經', 'https://zh.cn.nikkei.com/industry/itelectric-appliance.html','dl', 'newsContent02', 'dt', '', 'a', '', 'a', False)
  basicCrawler('日經', 'https://zh.cn.nikkei.com/top/2019-08-28-07-59-02.html','dl', 'newsContent02', 'dt', '', 'a', '', 'a', False)
  basicCrawler('日經', 'https://zh.cn.nikkei.com/top/2019-08-29-01-32-09.html','dl', 'newsContent02', 'dt', '', 'a', '', 'a', False)
  
  try:
    dom = requests.get('http://www.taipeitimes.com/News/front',  verify = False)
    dom = dom.content.decode('utf-8')
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find_all(['h1','h2']):
        articles.append(Article.from_list(ele.find('a').text,'http://www.taipeitimes.com/' + ele.find('a').get('href'),'Taipei Times'))

    dom = requests.get('http://www.taipeitimes.com/News/taiwan',  verify = False)
    dom = dom.content.decode('utf-8')
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find_all(['h1','h2']):
        articles.append(Article.from_list(ele.find('a').text,'http://www.taipeitimes.com/' + ele.find('a').get('href'),'Taipei Times'))
  
    dom = requests.get('http://www.taipeitimes.com/News/biz',  verify = False)
    dom = dom.content.decode('utf-8')
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find_all(['h1','h2']):
        articles.append(Article.from_list(ele.find('a').text,'http://www.taipeitimes.com/' + ele.find('a').get('href'),'Taipei Times'))
  except Exception as e:
    logging.warning(e)
  
  for i in range(1,3):
    articleCrawler('Reuters','https://www.reuters.com/news/archive/businessnews?view=page&page=' + str(i) + '&pageSize=10','article', 'h3', 'story-title', 'a', True)
  
  for i in range(1,3):
    articleCrawler('Reuters','https://www.reuters.com/news/archive/marketsNews?view=page&page=' + str(i) + '&pageSize=10','article', 'h3', 'story-title', 'a', True)
  
  for i in range(1,3):
    articleCrawler('Reuters','https://www.reuters.com/news/archive/technologynews?view=page&page=' + str(i) + '&pageSize=10','article', 'h3', 'story-title', 'a', True)
  
  for i in range(1,3):
    articleCrawler('Reuters','https://cn.reuters.com/news/archive/CNTopGenNews?view=page&page=' + str(i) + '&pageSize=10','article', 'h3', 'story-title', 'a', True)
  
  for i in range(1,3):
    articleCrawler('Reuters','https://cn.reuters.com/news/archive/CNAnalysesNews?view=page&page=' + str(i) + '&pageSize=10','article', 'h3', 'story-title', 'a', True)
  
  basicCrawler('BusinessKorea','http://www.businesskorea.co.kr/','div', 'auto-article auto-d04', 'li', '', 'strong', 'user-point', 'a', True)
  
  for i in range(1,3):
    basicCrawler('BusinessKorea','http://www.businesskorea.co.kr/news/articleList.html?sc_sub_section_code=S2N5&page=' + str(i), 'section', 'article-list-content text-left', 'div', 'table-row', 'strong', '', 'a', True)
    
  for i in range(1,3):
    basicCrawler('BusinessKorea','http://www.businesskorea.co.kr/news/articleList.html?sc_sub_section_code=S2N3&page=' + str(i), 'section', 'article-list-content text-left', 'div', 'table-row', 'strong', '', 'a', True)
  
  for i in range(1,3):
    basicCrawler('BusinessKorea','http://www.businesskorea.co.kr/news/articleList.html?sc_sub_section_code=S2N4&page=' + str(i), 'section', 'article-list-content text-left', 'div', 'table-row', 'strong', '', 'a', True)
  
  try:
    for i in range(1,3):
      dom = requests.get('https://en.yna.co.kr/news/' + str(i),  verify = False).text
      soup = BeautifulSoup(dom, 'lxml')
      for ele in soup.find_all('h2','tit'):
        articles.append(Article.from_list(ele.find('a').text, 'http:' + ele.find('a').get('href'),'韓聯社'))
    
    for i in range(1,3):
      dom = requests.get('https://en.yna.co.kr/economy/economy/' + str(i),  verify = False).text
      soup = BeautifulSoup(dom, 'lxml')
      for ele in soup.find_all('h2','tit'):
        articles.append(Article.from_list(ele.find('a').text, 'http:' + ele.find('a').get('href'),'韓聯社'))
          
    for i in range(1,3):
      dom = requests.get('https://en.yna.co.kr/market/finance/' + str(i),  verify = False).text
      soup = BeautifulSoup(dom, 'lxml')
      for ele in soup.find_all('h2','tit'):
        articles.append(Article.from_list(ele.find('a').text, 'http:' + ele.find('a').get('href'),'韓聯社'))
          
    for i in range(1,3):
      dom = requests.get('https://en.yna.co.kr/market/stocks/' + str(i),  verify = False).text
      soup = BeautifulSoup(dom, 'lxml')
      for ele in soup.find_all('h2','tit'):
        articles.append(Article.from_list(ele.find('a').text, 'http:' + ele.find('a').get('href'),'韓聯社'))
  except Exception as e:
    logging.warning(e)
  
  for i in range(1,3):
    basicCrawler('中時電子報','https://www.chinatimes.com/realtimenews/?page=' + str(i), 'ul', 'vertical-list list-style-none', 'li', '', 'h3', 'title', 'a', True)
  
  for i in range(1,3):
    basicCrawler('中時電子報','https://www.chinatimes.com/money/total?page=' + str(i), 'ul', 'vertical-list list-style-none', 'li', '', 'h3', 'title', 'a', True)
  
  for i in range(1,3):
    basicCrawler('中時電子報','https://www.chinatimes.com/chinese/?page=' + str(i), 'ul', 'vertical-list list-style-none', 'li', '', 'h3', 'title', 'a', True)
  
  for i in range(1,3):
    basicCrawler('中時電子報','https://www.chinatimes.com/technologynews/?page=' + str(i), 'ul', 'vertical-list list-style-none', 'li', '', 'h3', 'title', 'a', True)
  
  try:
    url = 'https://news.ltn.com.tw/list/breakingnews/all/'
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find('ul', 'list').find_all('li',''):
      if ele.find('p'):
        articles.append(Article.from_list(ele.find('p').text,  ele.find('a').get('href'),'自由時報'))
            
    url = 'https://news.ltn.com.tw/list/breakingnews/world/'
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find('ul', 'list').find_all('li',''):
      if ele.find('p'):
        articles.append(Article.from_list(ele.find('p').text,  ele.find('a').get('href'),'自由時報'))

    url = 'https://ec.ltn.com.tw/list/strategy/'       
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find('ul', 'list').find_all('li',''):
      if ele.find('p'):
        articles.append(Article.from_list(ele.find('p').text,  ele.find('a').get('href'),'自由時報'))

    url = 'https://ec.ltn.com.tw/list/international/'       
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find('ul', 'list').find_all('li',''):
      if ele.find('p'):
        articles.append(Article.from_list(ele.find('p').text,  ele.find('a').get('href'),'自由時報'))

    url = 'https://ec.ltn.com.tw/list/securities/'       
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find('ul', 'list').find_all('li',''):
      if ele.find('p'):
        articles.append(Article.from_list(ele.find('p').text,  ele.find('a').get('href'),'自由時報'))
  except Exception as e:
    logging.warning(url + ':' + str(e))
  
  for i in range(1,3):
    basicCrawler('MoneyDJ','https://www.moneydj.com/KMDJ/News/NewsRealList.aspx?a=MB010000&index1=' + str(i),'table','forumgrid', 'tr', '','a','', 'a', True)
  
  for i in range(1,3):
    basicCrawler('MoneyDJ','https://www.moneydj.com/KMDJ/News/NewsRealList.aspx?a=MB020000&index1=' + str(i),'table','forumgrid', 'tr', '','a','', 'a', True)

  for i in range(1,3):
    basicCrawler('MoneyDJ','https://www.moneydj.com/KMDJ/News/NewsRealList.aspx?a=MB070100&index1=' + str(i),'table','forumgrid', 'tr', '','a','', 'a', True)
  
  #動態載入
  basicCrawler('ETToday','https://www.ettoday.net/news/news-list.htm','div', 'part_list_2', 'h3', '', 'a', '', 'a', True)
  x = datetime.datetime.now()
  basicCrawler('ETToday','https://www.ettoday.net/news/news-list-' + x.strftime("%Y-%m-%d") + '-17.htm','div', 'part_list_2', 'h3', '', 'a', '', 'a', True)
  basicCrawler('ETToday','https://www.ettoday.net/news/news-list-' + x.strftime("%Y-%m-%d") + '-2.htm','div', 'part_list_2', 'h3', '', 'a', '', 'a', True)
  formattedCrawler('ETToday','https://www.ettoday.net/news_search/doSearch.php?keywords=%E5%8D%97%E4%BA%9E%E7%A7%91%E6%8A%80&idx=1','div','box_2','a','','a',False)
  
  try:
      dom = requests.get('https://tw.stock.yahoo.com/q/h?s=2408',  verify = False).text
      soup = BeautifulSoup(dom, 'lxml')
      for ele in soup.find('tbody').find_all('tbody')[1].find('tbody').find('tbody').find_all('tr'):
        if ele.find('a'):
          articles.append(Article.from_list(ele.find('a').text, 'https://tw.stock.yahoo.com' + ele.find('a').get('href'),'Yahoo奇摩股市'))
  except Exception as e:
    logging.warning('Yahoo奇摩股市:' + str(e))
  
  for i in range(1,3):
    articleCrawler('財訊科技','https://www.wealth.com.tw/home/articles?category_id=11&page=' + str(i),'article', 'h2', '', 'a', True)
  
  #惰性載入
  try:
    url = 'https://news.cnyes.com/news/cat/headline'
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find('div', '_2bFl theme-list').find_all('a', '_1Zdp'):
      articles.append(Article.from_list(ele['title'], url[:url.find('/', 8)] + ele.get('href'), '鉅亨網'))

    url = 'https://news.cnyes.com/trending'
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    for col in soup.find_all('nav', '_2QbU theme-trending-list'):
      for ele in col.find_all('a'):
        articles.append(Article.from_list(ele.get('title'), url[:url.find('/', 8)] + ele.get('href'), '鉅亨網'))

    url = 'https://news.cnyes.com/news/cat/tw_stock'
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find('div', '_2bFl theme-list').find_all('a', '_1Zdp'):
      articles.append(Article.from_list(ele['title'], url[:url.find('/', 8)] + ele.get('href'), '鉅亨網'))

    url = 'https://news.cnyes.com/news/cat/wd_stock'
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find('div', '_2bFl theme-list').find_all('a', '_1Zdp'):
      articles.append(Article.from_list(ele['title'], url[:url.find('/', 8)] + ele.get('href'), '鉅亨網'))
  except Exception as e:
    logging.warning('鉅亨網:' + str(e))
  
  #按鈕載入
  try:
    url = 'https://www.bnext.com.tw/articles'
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find('div', 'list_sty01').find_all('div','item_text_box'):
      articles.append(Article.from_list( ele.find('h2').text, ele.find('h2').parent.get('href'), '數位時代'))
  
    url = 'https://www.bnext.com.tw/categories/semiconductor'
    dom = requests.get(url,  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find('div', 'list_sty01').find_all('div','item_text_box'):
      articles.append(Article.from_list( ele.find('h2').text, ele.find('h2').parent.get('href'), '數位時代'))
  except Exception as e:
    logging.warning('數位時代:' + str(e))
  
  #按鈕載入
  basicCrawler('中央社','https://www.cna.com.tw/list/aall.aspx','ul', 'mainList imgModule', 'li', '','h2', '', 'a', False)
  basicCrawler('中央社','https://www.cna.com.tw/list/aie.aspx','ul', 'mainList imgModule', 'li', '','h2', '', 'a', False)  
  basicCrawler('中央社','https://www.cna.com.tw/list/acn.aspx','div', 'statement', 'li', '','h2', '', 'a', False)  
  basicCrawler('中央社','https://www.cna.com.tw/list/ait.aspx','ul', 'mainList imgModule', 'li', '','h2', '', 'a', False)
  
  for i in range(1,3):
    formattedCrawler('非凡', 'https://www.ustv.com.tw/UstvMedia/news/109?page=' + str(i), 'div', 'news_list', 'a', '', 'a', False)
  for i in range(1,3):
    formattedCrawler('非凡', 'https://www.ustv.com.tw/UstvMedia/news/103?page=' + str(i), 'div', 'news_list', 'a', '', 'a', False)
  
  #惰性載入
  articleCrawler('TechOrange','https://buzzorange.com/techorange/','article', 'h4', 'entry-title', 'a', False)
  
  formattedCrawler('SIA','https://www.semiconductors.org/news-events/latest-news/','div', 'resource-item', 'h3', '', 'a', False)
  
  basicCrawler('WSTS','https://www.wsts.org/76/PRESS-ARCHIVE','div', 'content-view-children', 'div', 'NewsArchive', 'a', '', 'a', True)
  
  basicCrawler('Trend Force Pressroom','https://press.trendforce.com.tw/press/','ul', 'list', 'li', '','a', '', 'a', False)
  
  try:
    dom = requests.get('http://www.icinsights.com/news/',  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find('div','content').find('div').find_all('a'):
      if ele.text != 'Join Our Mail List':
        articles.append(Article.from_list(ele.text, 'http://www.icinsights.com' + ele.get('href'), 'IC Insights Pressroom'))
  except Exception as e:
    logging.warning('IC Insights Pressroom:' + str(e))

  basicCrawler('SEMI Pressroom','http://www1.semi.org/zh/all-press-releases','div', 'main-container container', 'div', 'views-field views-field-title','span', 'greentext', 'a', True, 10)
  
  basicCrawler('SEMI Pressroom US','https://www.semi.org/en/node/7801','div', 'view-content', 'tr', '','a', '', 'a', True, 10)
  
  try:
    dom = requests.get('https://www.gartner.com/en/newsroom/archive',  verify = False).text
    soup = BeautifulSoup(dom, 'lxml')
    count = 0
    for ele in soup.find('div', 'newsroom-items').find_all('a'):
      articles.append(Article.from_list(ele.text, 'https://www.gartner.com/' + ele.get('href'), 'Gartner Pressroom'))
      count += 1
      if count >= 10:
        break
  except Exception as e:
    logging.warning("Gartner Pressroom:" + str(e))
  
  basicCrawler('Samsung Pressroom','https://news.samsung.com/global/category/press-resources/','ul', 'item', 'li', '', 'span', 'title ellipsis', 'a', False, 10)
 
  articleCrawler('SK Hynix Pressroom','https://news.skhynix.com/press-center/?blog_type=list_style','article', 'h3', 'title', 'a', False)
  
  try:
      dom = requests.get('http://investors.micron.com/press-releases',  verify = False).text
      soup = BeautifulSoup(dom, 'lxml')
      count = 0
      for ele in soup.find('div','nir-widget--list pressWrap').find_all('div','pressTitle'):
        articles.append(Article.from_list(ele.find_all('a')[1].text, 'http://investors.micron.com' + ele.find_all('a')[1].get('href'), 'Micron Pressroom'))
        count += 1
        if count >= 10:
          break
  except Exception as e:
    logging.warning('Micron Pressroom:' + str(e))

  try:
    for i in range(1,5):
      dom = requests.get('https://sina.com.hk/p/news/main/index/2/70/finance/realtime?p=' + str(i),  verify = False).text
      soup = BeautifulSoup(dom, 'lxml')
      for ele in soup.find('div', 'newsListing default finance').find_all('div', {'block noImage','block oneImage'}):
        articles.append(Article.from_list(ele.find('a').text.strip(), 'https://sina.com.hk/' + ele.find('a').get('href'), '新浪香港'))
  except Exception as e:
    logging.warning('新浪香港:' +str(e))
  
  try:
    dom = requests.get('http://economics.dwnews.com/big5/index.html',  verify = False)
    dom = dom.content.decode('utf-8')
    soup = BeautifulSoup(dom, 'lxml')
    for ele in soup.find('div', 'eco-l').find_all('li'):
      articles.append(Article.from_list(ele.find('div', 'mb').text.strip(), ele.find('a').get('href'), '多維新聞網'))
  except Exception as e:
    logging.warning('多維新聞網:' +str(e))

  # formattedCrawler('騰訊新聞', 'https://new.qq.com/ch/tech/','div', 'lazyload-placeholder', 'div', '', 'a', False)
  # basicCrawler('PTT','https://www.ptt.cc/bbs/Stock/index.html','div','r-list-container action-bar-margin bbs-screen','div','title','a','','a',True)

def outputWord():
  x = datetime.datetime.now()

  document = Document(r'Resources\Template.docx')

  document.styles['Normal'].font.name = u'微軟正黑體'
  document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微軟正黑體')

  table = document.tables[0]
  
  paragraph = table.cell(0,0).paragraphs[1]
  #run = p.add_run(x.strftime("%Y.%m.%d"))
  paragraph.text = x.strftime("%Y.%m.%d")
  run = paragraph.runs
  run[0].font.color.rgb = RGBColor(68, 84, 106)  # 顏色設置，這裏是用RGB顏色
  run[0].font.size = Pt(18)  # 字體大小設置，和word裏面的字號相對應
  #run[0].font.name = 'Microsoft JhengHei'

  # row = table.add_row()
  # #row.height = Inches(0.2)
  # row.cells[0].text = '南亞科相關新聞'
  # paragraph = row.cells[0].paragraphs[0]
  # run = paragraph.runs
  # run[0].font.color.rgb = RGBColor(255, 255, 255)  # 顏色設置，這裏是用RGB顏色
  # run[0].font.size = Pt(12)  # 字體大小設置，和word裏面的字號相對應
  # run[0].font.name = 'Microsoft JhengHei'
  # paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
  # shading_elm = parse_xml(r'<w:shd {} w:fill="767171"/>'.format(nsdecls('w')))
  # row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)

  
  # cell_1 = row.cells[0]
  # cell_2 = row.cells[2]
  # cell_1.merge(cell_2)

  digiTimesCount = 0
  i = 0
  for result in results:
    firstRun = False

    if i == 0 and len(result) != 0:
      firstRun = True
    elif i == 1 and len(result) != 0:
      row = table.add_row()
      #row.height = Inches(0.2)
      row.cells[0].text = '記憶體產業相關新聞'
      paragraph = row.cells[0].paragraphs[0]

      
      #paragraph.add_run( '記憶體相關新聞')
      #run = paragraph.runs
      #for run in paragraph.runs:
      paragraph.runs[0].font.color.rgb = RGBColor(255, 255, 255)  # 顏色設置，這裏是用RGB顏色
      paragraph.runs[0].font.size = Pt(12)  # 字體大小設置，和word裏面的字號相對應
      #paragraph.runs[0].font.name = 'Microsoft JhengHei'
      paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
      shading_elm = parse_xml(r'<w:shd {} w:fill="767171"/>'.format(nsdecls('w')))
      row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)
      paragraph_format = paragraph.paragraph_format
      paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
      paragraph_format.line_spacing = Pt(16)

      cell_1 = row.cells[0]
      cell_2 = row.cells[3]
      cell_1.merge(cell_2)
    elif i == 2 and len(result) != 0:
      row = table.add_row()
      #row.height = Inches(0.2)

      row.cells[0].text = '半導體產業相關新聞'
      paragraph = row.cells[0].paragraphs[0]
      #paragraph.add_run('半導體相關新聞')
      #run = paragraph.runs
      #for run in paragraph.runs:
      paragraph.runs[0].font.color.rgb = RGBColor(255, 255, 255)  # 顏色設置，這裏是用RGB顏色
      paragraph.runs[0].font.size = Pt(12)  # 字體大小設置，和word裏面的字號相對應
      #paragraph.runs[0].font.name = 'Microsoft JhengHei'
      paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
      shading_elm = parse_xml(r'<w:shd {} w:fill="767171"/>'.format(nsdecls('w')))
      row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)
      paragraph_format = paragraph.paragraph_format
      paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
      paragraph_format.line_spacing = Pt(16)

      cell_1 = row.cells[0]
      cell_2 = row.cells[3]
      cell_1.merge(cell_2)
    elif i == 3 and len(result) != 0:
      row = table.add_row()
      #row.height = Inches(0.1)
      row.cells[0].text = '市場狀況相關新聞'
      paragraph = row.cells[0].paragraphs[0]
      #for run in paragraph.runs:
      paragraph.runs[0].font.color.rgb = RGBColor(255, 255, 255)  # 顏色設置，這裏是用RGB顏色
      paragraph.runs[0].font.size = Pt(12)  # 字體大小設置，和word裏面的字號相對應
      #paragraph.runs[0].font.name = 'Microsoft JhengHei'
      paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
      shading_elm = parse_xml(r'<w:shd {} w:fill="767171"/>'.format(nsdecls('w')))
      row.cells[0]._tc.get_or_add_tcPr().append(shading_elm)
      paragraph_format = paragraph.paragraph_format
      paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
      paragraph_format.line_spacing = Pt(16)

      cell_1 = row.cells[0]
      cell_2 = row.cells[3]
      cell_1.merge(cell_2)
      #firstRun = True

    i += 1
    for article in result:
        if article.text == '':
          contentCrawlerView(article)

        if firstRun:
          row = table.rows[2]
          firstRun = False
        else:
          row = table.add_row()

        p = row.cells[1].paragraphs[0]

        # paragraph_format = document.styles['Normal'].paragraph_format
        # paragraph_format.space_before = 0
        # p.paragraph_format.line_spacing = 0.5
        

        #p = document.add_paragraph('')
        #if article.source == 'DIGITIMES電子時報':
        add_hyperlink(p, article.title, 'Article-' + str(digiTimesCount), False)
        digiTimesCount += 1
        #else:
        #  add_hyperlink(p, article.title, article.url)

        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(16)

        p = row.cells[1].add_paragraph(article.source)
        if article.author != '':
          p.add_run('\t' + article.author)
        if article.date != '':
          p.add_run('\t' + article.date)
        # for run in paragraph.runs:
        #   run.font.name = 'Microsoft JhengHei'
        # for run in paragraph.runs:
        #   #run[0].font.color.rgb = RGBColor(255, 255, 255)  # 顏色設置，這裏是用RGB顏色
        #   run.font.size = Pt(12)  # 字體大小設置，和word裏面的字號相對應
        #   run.font.name = 'Microsoft JhengHei'
        
        try:
          if article.img_name is not None and article.img_name != '' and os.path.isfile(os.path.join('Temp', article.img_name)) :
            paragraph = row.cells[0].paragraphs[0]
            run = paragraph.add_run()
            paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER

            row.cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            #run.add_picture('Nanya.jpg', width=Inches(1))
            #title = article.img_url[article.img_url.rfind('/')+1:]
            #if title.endswith('jpg') == False and title.endswith('png') == False and title.endswith('gif') == False:
            #  title += '.jpg'
            #row.cells[0].width = Inches(2)
            run.add_picture(os.path.join('Temp', article.img_name), width=Inches(2))
            #os.remove('Temp\'' + title)
        except Exception as e:
          logging.warning('Output Word Image Error:' + article.title + str(e))
        
        artparagraphs = article.text.splitlines()
        #logging.warning(len(artparagraphs))
        
        text = ''
        num = 0

        while len(text) < 150 and num < len(artparagraphs):
          text += artparagraphs[num]
          num +=1

        #paragraph = row.cells[1].add_paragraph(text)
        p.add_run('\n' + text)
        #paragraph.line_spacing_rule = WD_LINE_SPACING.SINGLE
        #run = paragraph.runs
        for run in paragraph.runs:
          #run[0].font.color.rgb = RGBColor(255, 255, 255)  # 顏色設置，這裏是用RGB顏色
          run.font.size = Pt(12)  # 字體大小設置，和word裏面的字號相對應
          run.font.name = 'Microsoft JhengHei'

        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(16)
        
        cell_1 = row.cells[1]
        cell_2 = row.cells[2]
        cell_1.merge(cell_2)

        # #logging.warning(article.text)
        # document.add_paragraph('', style='Intense Quote')
        #i += 1

        for cell in row.cells:
          set_cell_border(
            cell,
            bottom={"sz": 5, "color": "#767171", "val": "single"}
          )

  if len(results[0])==0:
    #row = table.rows[n]
    remove_row(table, table.rows[1])
    remove_row(table, table.rows[1])
   
  #modifyBorder(table)

  # document.add_page_break()
  paragraphIndex = 3
  tableIndex = 1
 
  digiTimesCount = 0
  #table.cell(0,0).text = 'DIGITIMES NEWS'
  for result in results:
    for article in result:
      #if article.source == 'DIGITIMES電子時報':
        if tableIndex != 1:
          copy_table_after(table, document.paragraphs[4])
          # paragraph = document.paragraphs[4]
          # run = paragraph.add_run()
          # run.add_break(WD_BREAK.PAGE)

        table = document.tables[tableIndex]
        #table = document.add_table(rows=0, cols=1)

        tableIndex += 1
        paragraphIndex +=1

        row = table.rows[0]
        p = row.cells[0].paragraphs[0]

        #p = document.add_paragraph('')
        p.text = ''
        p.add_run(article.title)

        #run = p.runs
        add_bookmark(p.runs[0],'Article-' + str(digiTimesCount))
        digiTimesCount += 1
        p.runs[0].font.size = Pt(16)  # 字體大小設置，和word裏面的字號相對應
        #run[0].font.name = 'Microsoft JhengHei'
        p.runs[0].font.bold = True

        for i in range(len(row.cells[0].paragraphs)-2):
          delete_paragraph(row.cells[0].paragraphs[1])
        
        p = insert_paragraph_after(p)
        p.add_run(article.source)
        #p = row.cells[0].add_paragraph(article.source)
        if article.author != '':
          p.add_run('\t' + article.author)
        if article.date != '':
          p.add_run('\t' + article.date)
        # for run in paragraph.runs:
        #   run.font.name = 'Microsoft JhengHei'

        # paragraph = row.cells[0].add_paragraph(article.text)
        p.add_run('\n' + article.text)

        paragraph_format = p.paragraph_format
        paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph_format.line_spacing = Pt(16)
        #add_hyperlink(p, '<回到電子報>', '')

        
        
        # for run in paragraph.runs:
        #   #run[0].font.color.rgb = RGBColor(255, 255, 255)  # 顏色設置，這裏是用RGB顏色
        #   run.font.size = Pt(12)  # 字體大小設置，和word裏面的字號相對應
        #   run.font.name = 'Microsoft JhengHei'
        #   run.font.bold = True
        # #document.add_page_break()

  try:
    document.save('Memory Related News ' + x.strftime("%m%d") + '.docx')
  except Exception as e:
    logging.warning('outputWord Output Word Error:' + str(e))

  folder = 'Temp'
  for the_file in os.listdir(folder):
    file_path = os.path.join(folder, the_file)
    try:
        if os.path.isfile(file_path):
            os.unlink(file_path)
    except Exception as e:
        logging.warning(e)

def outputTxt():
  x = datetime.datetime.now()

  try:
    fp = open("Raw Titles.txt", 'w+', encoding='utf-8')
    #fp.write('News List\n')

    for article in articles:
      fp.write(article.title + '\n')
      #logging.warning(article.title)
    
    #fp.write('----------------\n')

    fp.close()
  except Exception as e:
    logging.warning(e)

  try:
    fp = open("News List.txt", "w+")
    #fp.write('Raw titles\n')

    for url in archive:
      fp.write(url + '\n')

    for result in results:
      #logging.warning(len(result))
      for article in result:
        fp.write(article.url.strip() + '\n')
          #logging.warning(article.title)

    fp.close()
  except Exception as e:
    logging.warning('outputWord News List Error:' + str(e))

  try:
    fp = open('Memory Related News ' + x.strftime("%m%d") + '-old.txt', 'w+', encoding='utf-8')
    fp.write('News List\n')
    
    for result in results:
      for article in result:
        fp.write(article.title + '\n')
        #logging.warning(article.title)
      
    fp.write('----------------\n')
    
    i=0
    for result in results:
      if i == 0 and len(result) != 0:
        #fp.write('----------------\n')
        fp.write('南亞科相關新聞\n')
        fp.write('----------------\n')
      elif i == 1 and len(result) != 0:
        #fp.write('----------------\n')
        fp.write('記憶體相關新聞\n')
        fp.write('----------------\n') 
      elif i == 2 and len(result) != 0:
        #fp.write('----------------\n')
        fp.write('半導體相關新聞\n')
        fp.write('----------------\n')
      elif i == 3 and len(result) != 0:
        #fp.write('----------------\n')
        fp.write('產業別相關新聞\n')
        fp.write('----------------\n')

      i += 1   
      for article in result: 
          #logging.warning(article.text)
          if article.text == '':
            contentCrawlerView(article)

          try:
            #filtered = filter(lambda x: not re.match(r'^\s*$', x), article.text)
            #logging.warning(article.author+'\n')
            #logging.warning(article.date+'\n')
            fp.write(article.url+'\n')
            fp.write(article.title+'\n')
            fp.write(article.author+'\n')
            fp.write(article.date+'\n')
            fp.write(article.source+'\n')
            fp.write(article.text +'\n')
            fp.write('----------------\n')
          except Exception as e:
            logging.warning(e)

    fp.close()
  except Exception as e:
    logging.warning(e)

  try:
    fp = open('Memory Related News ' + x.strftime("%m%d") + '.txt', 'w+', encoding='utf-8')
    fp.write('<NewsList>\n')
    #fp.write('News List\n')
    
    # for result in results:
    #   for article in result:
    #     fp.write(article.title + '\n')
    #     #logging.warning(article.title)
      
    #fp.write('----------------\n')
    # first = True
    
    for result in results:  
      for article in result:
          # if first == False:
          #   fp.write(',\n')
          
          # first = False
          fp.write('<News>\n')
          #logging.warning(article.text)
          if article.text == '':
            contentCrawlerView(article)

          try:
            #filtered = filter(lambda x: not re.match(r'^\s*$', x), article.text)
            #logging.warning(article.author+'\n')
            #logging.warning(article.date+'\n')
            fp.write('<URL>')
            fp.write(article.url+'</URL>\n')
            fp.write('<Title>')
            fp.write(article.title+'</Title>\n')
            fp.write('<Author>')
            fp.write(article.author.strip() + '\t' + article.date+'</Author>\n')
            fp.write('<Source>')
            fp.write(article.source+'</Source>\n')
            fp.write('<Text>')
            fp.write(article.text +'</Text>\n')
            fp.write('</News>\n')
          except Exception as e:
            logging.warning(e)

    fp.write('</NewsList>')
    fp.close()
  except Exception as e:
    logging.warning(e)

def sendMail():
  x = datetime.datetime.now()

  sender = 'lcw726@ntc.com.tw'
  #receivers = config['Mail']['receivers']
  receivers = config['receivers']
  #receivers = ('lcw726@ntc.com.tw','lcw726@ntc.com.tw') #('zehui@ntc.com.tw','joyinhung@ntc.com.tw','sandraliu@ntc.com.tw','lcw726@ntc.com.tw') #
 
  # 創建發送郵件的類別物件
  message = MIMEMultipart()
  message['From'] = sender
  message['To'] = ", ".join(receivers)
  subject = 'Memory Related News ' + x.strftime("%m%d")
  message['Subject'] = Header(subject, 'utf-8')
  
  # 郵件正文內容
  message.attach(MIMEText('Hi all,\n\n附件為今日新聞電子報檔案，請過目，謝謝。', 'plain', 'utf-8'))
  
  # 附件1
  att1 = MIMEText(open('Memory Related News ' + x.strftime("%m%d") + '.docx', 'rb').read(), 'base64', 'utf-8')
  att1["Content-Type"] = 'application/octet-stream'
  # 可更改附件名稱
  att1["Content-Disposition"] = 'attachment; filename="Memory Related News ' + x.strftime("%m%d") + '.docx"'
  message.attach(att1)

  # 附件2
  att2 = MIMEText(open('Memory Related News ' + x.strftime("%m%d") + '.txt', 'rb').read(), 'base64', 'utf-8')
  att2["Content-Type"] = 'application/octet-stream'
  # 可更改附件名稱
  att2["Content-Disposition"] = 'attachment; filename="Memory Related News ' + x.strftime("%m%d") + '.txt"'
  message.attach(att2)

  # 附件3
  att3 = MIMEText(open('Raw Titles.txt', 'rb').read(), 'base64', 'utf-8')
  att3["Content-Type"] = 'application/octet-stream'
  # 可更改附件名稱
  att3["Content-Disposition"] = 'attachment; filename="Raw Titles.txt"'
  message.attach(att3)

  # 附件4
  att4 = MIMEText(open('Memory Related News ' + x.strftime("%m%d") + '-old.txt', 'rb').read(), 'base64', 'utf-8')
  att4["Content-Type"] = 'application/octet-stream'
  # 可更改附件名稱
  att4["Content-Disposition"] = 'attachment; filename="Memory Related News ' + x.strftime("%m%d") + '-old.txt"'
  message.attach(att4)

  # 附件5
  att5 = MIMEText(open('News List.txt', 'rb').read(), 'base64', 'utf-8')
  att5["Content-Type"] = 'application/octet-stream'
  # 可更改附件名稱
  att5["Content-Disposition"] = 'attachment; filename="News List.txt"'
  message.attach(att5)
  
  try:
    smtpObj = smtplib.SMTP("smtp.gmail.com", 587)
    smtpObj.ehlo()
    smtpObj.starttls()
    smtpObj.login("lcw726@gmail.com", "google24det")
    smtpObj.sendmail(sender, receivers, message.as_string())
    logging.info("郵件發送成功")
  except smtplib.SMTPException as e:
      logging.warning("Error: 無法發送郵件:" + str(e))

def deleteFiles():
  x = datetime.datetime.now()
  try:
      os.unlink('Memory Related News ' + x.strftime("%m%d") + '.docx')
      os.unlink('Memory Related News ' + x.strftime("%m%d") + '.txt')
      os.unlink('Memory Related News ' + x.strftime("%m%d") + '-old.txt')
      os.unlink('Raw Titles.txt')
  except Exception as e:
      logging.warning(e)  
 
if __name__ == '__main__':
  requests.packages.urllib3.disable_warnings()
  logging.basicConfig(filename='app.log', filemode='a', format='%(asctime)s %(message)s', datefmt='%m/%d/%Y %I:%M:%S %p')

  getNewsList()
  filterArticles()
  outputWord()
  outputTxt()
  sendMail()
  deleteFiles()